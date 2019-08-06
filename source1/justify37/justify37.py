import win32com.client as win32
import os 
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import re
from win32com.client import constants
##Read the path 
l=[]
count=0
import sys                                                                                 
iter=sys.argv[1]
start=datetime.now()
print("-----------------------------------------------------------------------------------------------------------------")
print("Document Name:", iter)
print("CheckList Rule - 32: Checking 'Justify' alignment for paragraphs.")
print("Document Review Start Time:", start,"HH:MM:SS")
print("-----------------------------------------------------------------------------------------------------------------")
print("\n")

##Open the Document
start=datetime.now()
if iter.endswith('.doc') :
 word1 = win32.gencache.EnsureDispatch ("Word.Application")
 word1.Visible = True
 p = os.path.abspath(iter)
 word1.Documents.Open(p)
 sheet_1 = word.ActiveDocument
 for i in range(1,sheet_1.Paragraphs.Count+1):
  t=sheet_1.Paragraphs(i).Range.Text.encode('ascii','ignore').decode()
  t=t.strip('\r')
  t=t.strip('\r\x07')
  t=t.strip('\x0c')
  t=t.strip('\x0b')
  t=t.strip('\x0a')
  t=t.rstrip(' ')
  t=t.strip('\n')
  n=sheet_1.Paragraphs(i).Range.Style
  if re.search("Heading",str(n)):
   l.append(t)
  if re.search("Normal",str(n)) and sheet_1.Paragraphs(i).Alignment != win32.constants.wdAlignParagraphJustify and str(t).strip()!='':
   print("String:",(sheet_1.Paragraphs(i)).Range.Text.encode('ascii','ignore').decode())
   print("Length:",len(t))
   print("Page number:",sheet_1.Paragraphs(i).Range.Information(constants.wdActiveEndAdjustedPageNumber))
   print("Line On Page:",sheet_1.Paragraphs(i).Range.Information(constants.wdFirstCharacterLineNumber))
   print("\n")
   #l.append((doc.Paragraphs(i)).Range.Text.encode('ascii','ignore').decode())
   count=count+1
 sheet_1.Close()
 word1.Quit()
elif iter.endswith('.docx'):
 doc = Document(iter)
 para=list(doc.paragraphs)
 for i in range(len(para)):
  t=doc.paragraphs[i].text.encode('ascii','ignore').decode()
  if (len(doc.paragraphs[i].text))!=0:
   n=doc.paragraphs[i].style.name.encode('ascii','ignore').decode()
   #print(n)
   if re.search("Heading",n):
    l.append(doc.paragraphs[i].text.encode('ascii','ignore').decode())
   if re.search("Normal",n) and doc.paragraphs[i].alignment != WD_ALIGN_PARAGRAPH.JUSTIFY and str(t).strip()!='':
    if l!=[]:
     print("\n")
     print("Heading Section:",l[-1])
     print("String:",doc.paragraphs[i].text.encode('ascii','ignore').decode())
     print("\n")
    else:
     #print("Heading Section:",l[-1])
     print("String:",doc.paragraphs[i].text.encode('ascii','ignore').decode())
     print("\n")
    #l.append(doc.paragraphs[i].text.encode('ascii','ignore').decode())
    count=count+1
if count>0:
 print("Status:Fail")
else:
 print("Status:Pass")
end=datetime.now()
print("\nDocument Review End Time:", end)
print("\nTime taken For Document Review:", end-start,"HH:MM:SS")  