import win32com.client as win32
import os 
from datetime import datetime
import string
import re
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from win32com.client import constants
count=0
iter=sys.argv[1]
start=datetime.now()
print("--------------------------------------------------------------------------------------------------------")
print("Document Name:", iter)
print("CheckList Rule - 14: Check the Caption for figures and Table is center.")
print("Document Review Start Time:", start,"HH:MM:SS")
print("--------------------------------------------------------------------------------------------------------")
print("\n")

l=[]
ll=[]
lll=[]
##Open the Document
if iter.endswith('.doc'):
 word1 = win32.Dispatch("Word.Application")
 word1.Visible = True
 p = os.path.abspath(iter)
 word1.Documents.Open(p)
 sheet_1 = word1.ActiveDocument
 par=sheet_1.Paragraphs
 try:
  for para in par:
   #m=str(para)
   #m=m.rstrip('\r\x07')
   #m=m.rstrip('\f')
   #t=str.maketrans({key: None for key in string.punctuation})
   #k=m.translate(t)
   #j=k.encode("utf-8")
   j=para.Range.Text.encode('ascii','ignore').decode()
   if len(j)!=0 and j!='':
    a=para.Range.Style
    #print(a)
    if len(j)>0 and str(a)=="Caption"and para.Alignment != win32.constants.wdAlignParagraphCenter:
     if re.search("[\t\s]+?Table",j) or re.search("^Table",j):
      print("String:",str(para))
      print("Page number:",para.Range.Information(constants.wdActiveEndAdjustedPageNumber))
      print("Line On Page:",para.Range.Information(constants.wdFirstCharacterLineNumber))
      count=count+1
      print("-----------------------------------------------------------------------------------------------------------------------------")
     if re.search("[\t\s]+?Figure",j) or re.search("^Figure",j):
      print("String:",str(para))
      print("Page number:",para.Range.Information(constants.wdActiveEndAdjustedPageNumber))
      print("Line On Page:",para.Range.Information(constants.wdFirstCharacterLineNumber))
      count=count+1
      print("-----------------------------------------------------------------------------------------------------------------------------")
  if count>0:
   print("Status:Fail")
  else:
   print("Status:Pass")
   
 except:
  pass
 sheet_1.Close()
 word1.Quit()
elif iter.endswith('.docx'):
 doc = Document(iter)
 para=list(doc.paragraphs)
 for i in range(len(para)):
  p=doc.paragraphs[i].text.encode('ascii','ignore').decode()
  #print(p)
  if len(p)>0 and (str(doc.paragraphs[i].style.name)=="Caption") and (doc.paragraphs[i].alignment != WD_ALIGN_PARAGRAPH.CENTER):#and ((re.search('^Figure',doc.paragraphs[i].text)) or (re.search('^Table',doc.paragraphs[i].text))) and (doc.paragraphs[i].alignment != WD_ALIGN_PARAGRAPH.CENTER):
   if re.search("[\t\s]+?Table",p) or re.search("^Table",p):
    print("String:",doc.paragraphs[i].text.encode('ascii','ignore').decode()) 
    count=count+1
   if re.search("[\t\s]+?Figure",p) or re.search("^Figure",p):
    print("String:",doc.paragraphs[i].text.encode('ascii','ignore').decode()) 
    count=count+1
 if count>0:
  print("\nStatus:Fail")
 else:
  print("\nStatus:Pass")
print(datetime.now()-start,"HH:MM:SS")

