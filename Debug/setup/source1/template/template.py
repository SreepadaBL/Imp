import win32com.client as win32
import os 
import re
import sys
from win32com.client import constants
from docx import Document
from docx.enum.text import WD_BREAK
from datetime import datetime
l=[]
ll=[]
lll=[]
llll=[]
import sys                                                                                 
iter=sys.argv[1]
start=datetime.now()
count=0
flag=0
print("-----------------------------------------------------------------------------------------------------------------")
print("Document Name:", iter)
print("CheckList Rule - 6: Document Template Check for SRS, SDS, User Manual, Release Notes, HLD, LLD, Test plan, SOW. ")
print("Document Review Start Time:", start,"HH:MM:SS")
print("-----------------------------------------------------------------------------------------------------------------")
print("\n")
##Open the Document
if iter.endswith('.doc'):
 word1 = win32.Dispatch("Word.Application")
 word1.Visible = True
 p = os.path.abspath(iter)
 word1.Documents.Open(p)
 sheet_1 = word1.ActiveDocument
 title=sheet_1.BuiltInDocumentProperties('Title')
 if title!='':
  l.append(title)
  count=count+1
  for p in sheet_1.Paragraphs:
   a=p.Range.Style
   #k=a.encode('ascii','ignore').decode()
   #print(k)
   if re.search("Heading [0-9]",str(a)):
    c=p.Range.Text.encode('ascii','ignore').decode().lower()
    c=c.strip('\r')
    c=c.strip('\r\x07')
    c=c.strip('\x0c')
    c=c.strip('\x0b')
    c=c.strip('\x0a')
    c=c.rstrip(' ')
    c=c.strip('\n')
    ll.append(c)			##list containing all the strings with heading style
   
  sow=['introduction','project overview and scope','definitions and acronyms','references','assumptions, dependencies, constraints and risks','high level requirements/target specifications','estimation summary','responsibility and ownership','key milestones and work products to be delivered','release definition','development and test environment','knowledge transfer','verification methods and acceptance criteria','tracking and reporting','product delivery mechanism','attachments']
  srs=['introduction','assumptions and risks','dependencies, constraints and limitations','overall description','functional requirements','interface requirements','non-functional requirements','error handling','compiler defined options','acceptance criteria','appendix']
  sds=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design - (modules 1..n)','sub-module 1..n design','sub-module 1..n pseudo code','integration approach and sequence','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  um=['introduction','environment setup','limitations and constraints','over all description','states of the system','[module_name]','source code build procedure','appendix i','trouble shooting']
  rn=['project details','contents, media and release method','brief introduction of product','summary of release history','environment required','limitations and known issues','installation procedure','components of release']
  tp=['introduction','definitions and acronyms','references','assumptions, dependencies and constraints','scope','test environment','features not tested']
  hld=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design – (modules 1..n)','integration approach and sequence:','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  lld=['introduction','definitions and acronyms','references','conventions and standards','data definition','sub-module 1..n design','sub-module 1..n pseudo code']
  cmp=['scope','acronyms and definitions','cm tools','project environment','configuration identification','naming conventions','location of cis','versioning','baselining','configuration control','configuration status accounting','configuration management audit','backup','release mechanism']
  pp=['introduction','definitions and acronyms','strategy and approach','assumptions, constraints and dependencies','project organization structure','risk management plan','summary of estimates by phase / delivery milestones','development environment','project monitoring and control /status reporting/ communication plan','quantitative objectives, measurement and data management plan','quality management (verification, validation and causal analysis)','product release plan','standards to be followed','decision management plan','tailoring and deviations','reference sheet_1ument']

   
  if re.search("user manual",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in um:
     if re.search(j,i):
      lll.append(j)			##lll is list of headings found in main sheet_1ument and are present in table of contents.
   result=set(um)-set(lll)
  elif re.search("Release Notes",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in rn:
     if re.search(j,i):
      lll.append(j)
   result=set(rn)-set(lll)
  
  elif re.search("SRS|Software Requirement Specification|Software Requirements Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in srs:
     if re.search(j,i):
      lll.append(j)
   result=set(srs)-set(lll)
   
  elif re.search("Statement of Work|Scope of Work",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sow:
     if re.search(j,i):
      lll.append(j)
   result=set(sow)-set(lll)
   
  elif re.search("Project Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in pp:
     if re.search(j,i):
      lll.append(j)
   result=set(pp)-set(lll)
   
  elif re.search("^Test Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in tp:
     if re.search(j,i):
      lll.append(j)
   result=set(tp)-set(lll)
   
  elif re.search("Low Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in lld:
     if re.search(j,i):
      lll.append(j)
   result=set(lld)-set(lll)
   
  elif re.search("High Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in hld:
     if re.search(j,i):
      lll.append(j)
   result=set(hld)-set(lll)
   
  elif re.search("Configuration Management Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in cmp:
     if re.search(j,i):
      lll.append(j)
   result=set(cmp)-set(lll)
   
  elif re.search("SDS|Software Design Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sds:
     if re.search(j,i):
      lll.append(j)
   result=set(sds)-set(lll)
  if result==set():
   print("Document template is according to QMS.")
   print("Status:Pass")
  else:
   for x in result:
    print("Heading Not Found:","\t",x.upper()) 
   print("Status:Fail")    
 elif title1=='':
  for p in sheet_1.Paragraphs:
   a=p.style.name
   k=a.encode('ascii','ignore').decode()
   if re.search("Heading [0-9]",k):
    c=p.text.encode('ascii','ignore').decode().lower()
    c=c.strip('\r')
    c=c.strip('\r\x07')
    c=c.strip('\x0c')
    c=c.strip('\x0b')
    c=c.strip('\x0a')
    c=c.rstrip(' ')
    c=c.strip('\n')
    #print(p.text.encode('ascii','ignore').decode())
    ll.append(c)			##list containing all the strings with heading style
   d=p.text.encode('ascii','ignore').decode()
   if re.search('User Manual',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Statement of Work|Scope of Work',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Project Plan',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('SDS|Software Design Specification',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('SRS|Software Requirement Specification|Software Requirements Specification',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Test Plan',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Release Notes',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Low Level Design',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('High Level Design',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Configuration Management Plan',d,re.IGNORECASE):
    l.append(d)
   
 
  sow=['introduction','project overview and scope','definitions and acronyms','references','assumptions, dependencies, constraints and risks','high level requirements/target specifications','estimation summary','responsibility and ownership','key milestones and work products to be delivered','release definition','development and test environment','knowledge transfer','verification methods and acceptance criteria','tracking and reporting','product delivery mechanism','attachments']
  srs=['introduction','assumptions and risks','dependencies, constraints and limitations','overall description','functional requirements','interface requirements','non-functional requirements','error handling','compiler defined options','acceptance criteria','appendix']
  sds=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design - (modules 1..n)','sub-module 1..n design','sub-module 1..n pseudo code','integration approach and sequence','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  um=['introduction','environment setup','limitations and constraints','over all description','states of the system','[module_name]','source code build procedure','appendix i','trouble shooting']
  rn=['project details','contents, media and release method','brief introduction of product','summary of release history','environment required','limitations and known issues','installation procedure','components of release']
  tp=['introduction','definitions and acronyms','references','assumptions, dependencies and constraints','scope','test environment','features not tested']
  hld=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design – (modules 1..n)','integration approach and sequence:','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  lld=['introduction','definitions and acronyms','references','conventions and standards','data definition','sub-module 1..n design','sub-module 1..n pseudo code']
  cmp=['scope','acronyms and definitions','cm tools','project environment','configuration identification','naming conventions','location of cis','versioning','baselining','configuration control','configuration status accounting','configuration management audit','backup','release mechanism']
  pp=['introduction','definitions and acronyms','strategy and approach','assumptions, constraints and dependencies','project organization structure','risk management plan','summary of estimates by phase / delivery milestones','development environment','project monitoring and control /status reporting/ communication plan','quantitative objectives, measurement and data management plan','quality management (verification, validation and causal analysis)','product release plan','standards to be followed','decision management plan','tailoring and deviations','reference sheet_1ument']

   
  if re.search("user manual",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in um:
     if re.search(j,i):
      lll.append(j)			##lll is list of headings found in main sheet_1ument and are present in table of contents.
   result=set(um)-set(lll)
  elif re.search("Release Notes",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in rn:
     if re.search(j,i):
      lll.append(j)
   result=set(rn)-set(lll)
  elif re.search("SRS|Software Requirement Specification|Software Requirements Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in srs:
     if re.search(j,i):
      lll.append(j)
   
   result=set(srs)-set(lll)
  elif re.search("Statement of Work|Scope of Work",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sow:
     if re.search(j,i):
      lll.append(j)
   result=set(sow)-set(lll)
  elif re.search("Project Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in pp:
     if re.search(j,i):
      lll.append(j)
   result=set(pp)-set(lll)
  elif re.search("^Test Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in tp:
     if re.search(j,i):
      lll.append(j)
   result=set(tp)-set(lll)
  elif re.search("Low Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in lld:
     if re.search(j,i):
      lll.append(j)
   result=set(lld)-set(lll)
  elif re.search("High Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in hld:
     if re.search(j,i):
      lll.append(j)
   result=set(hld)-set(lll)
  elif re.search("Configuration Management Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in cmp:
     if re.search(j,i):
      lll.append(j)
   result=set(cmp)-set(lll)
  elif re.search("SDS|Software Design Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sds:
     if re.search(j,i):
      lll.append(j)
   result=set(sds)-set(lll)
  if result==set():
   print("Document template is according to QMS.")
   print("Status:Pass")
  else:
   for x in result:
    print("Heading Not Found:","\t",x.upper()) 
   print("Status:Fail")   

 end=datetime.now()
 print("\nDocument Review End Time:", end)
 print("\nTime taken For Document Review:", end-start,"HH:MM:SS")  
 sheet_1.Close()
 word1.Quit() 
elif iter.endswith('.docx'):
 sheet_1=Document(iter)
 a=sheet_1.core_properties
 title1=a.title
 if title1!='':
  l.append(title1)
  count=count+1
  for p in sheet_1.paragraphs:
   a=p.style.name
   k=a.encode('ascii','ignore').decode()
   if re.search("Heading [0-9]",k):
    c=p.text.encode('ascii','ignore').decode().lower()
    c=c.strip('\r')
    c=c.strip('\r\x07')
    c=c.strip('\x0c')
    c=c.strip('\x0b')
    c=c.strip('\x0a')
    c=c.rstrip(' ')
    c=c.strip('\n')
    ll.append(c)			##list containing all the strings with heading style
 
  sow=['introduction','project overview and scope','definitions and acronyms','references','assumptions, dependencies, constraints and risks','high level requirements/target specifications','estimation summary','responsibility and ownership','key milestones and work products to be delivered','release definition','development and test environment','knowledge transfer','verification methods and acceptance criteria','tracking and reporting','product delivery mechanism','attachments']
  srs=['introduction','assumptions and risks','dependencies, constraints and limitations','overall description','functional requirements','interface requirements','non-functional requirements','error handling','compiler defined options','acceptance criteria','appendix']
  sds=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design - (modules 1..n)','sub-module 1..n design','sub-module 1..n pseudo code','integration approach and sequence','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  um=['introduction','environment setup','limitations and constraints','over all description','states of the system','[module_name]','source code build procedure','appendix i','trouble shooting']
  rn=['project details','contents, media and release method','brief introduction of product (summary of features in current release and changes from previous release if any):','summary of release history','environment required','limitations and known issues','installation procedure','components of release']
  tp=['introduction','definitions and acronyms','references','assumptions, dependencies and constraints','scope','test environment','features not tested']
  hld=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design – (modules 1..n)','integration approach and sequence:','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  lld=['introduction','definitions and acronyms','references','conventions and standards','data definition','sub-module 1..n design','sub-module 1..n pseudo code']
  cmp=['scope','acronyms and definitions','cm tools','project environment','configuration identification','naming conventions','location of cis','versioning','baselining','configuration control','configuration status accounting','configuration management audit','backup','release mechanism']
  pp=['introduction','definitions and acronyms','strategy and approach','assumptions, constraints and dependencies','project organization structure','risk management plan','summary of estimates by phase / delivery milestones','development environment','project monitoring and control /status reporting/ communication plan','quantitative objectives, measurement and data management plan','quality management (verification, validation and causal analysis)','product release plan','standards to be followed','decision management plan','tailoring and deviations','reference sheet_1ument']
  result=[]
  if re.search("user manual",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in um:
     if re.search(j,i):
      lll.append(j)			##lll is list of headings found in main sheet_1ument and are present in table of contents.
   result=set(um)-set(lll)
  elif re.search("Release Notes",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in rn:
     if re.search(j,i):
      lll.append(j)
   result=set(rn)-set(lll)
  
  elif re.search("SRS|Software Requirement Specification|Software Requirements Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in srs:
     if re.search(j,i):
      lll.append(j)
   result=set(srs)-set(lll)
   
  elif re.search("Statement of Work|Scope of Work",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sow:
     if re.search(j,i):
      lll.append(j)
   result=set(sow)-set(lll)
   
  elif re.search("Project Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in pp:
     if re.search(j,i):
      lll.append(j)
   result=set(pp)-set(lll)
  elif re.search("^Test Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in tp:
     if re.search(j,i):
      lll.append(j)
   result=set(tp)-set(lll)
   
  elif re.search("Low Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in lld:
     if re.search(j,i):
      lll.append(j)
   result=set(lld)-set(lll)
   
  elif re.search("High Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in hld:
     if re.search(j,i):
      lll.append(j)
   result=set(hld)-set(lll)
   
  elif re.search("Configuration Management Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in cmp:
     if re.search(j,i):
      lll.append(j)
   result=set(cmp)-set(lll)
   
  elif re.search("SDS|Software Design Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sds:
     if re.search(j,i):
      lll.append(j)
   result=set(sds)-set(lll)
  #print(result)
  if result==set():
   print("Document template is according to QMS.")
   print("Status:Pass")
  else:
   for x in result:
    print("Heading Not Found:","\t",x.upper()) 
   print("Status:Fail")   
 
	  
 elif title1=='':
  for p in sheet_1.paragraphs:
   a=p.style.name
   k=a.encode('ascii','ignore').decode()
   if re.search("Heading [0-9]",k):
    c=p.text.encode('ascii','ignore').decode().lower()
    c=c.strip('\r')
    c=c.strip('\r\x07')
    c=c.strip('\x0c')
    c=c.strip('\x0b')
    c=c.strip('\x0a')
    c=c.rstrip(' ')
    c=c.strip('\n')
    #print(p.text.encode('ascii','ignore').decode())
    ll.append(c)			##list containing all the strings with heading style
   d=p.text.encode('ascii','ignore').decode()
   if re.search('User Manual',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Statement of Work|Scope of Work',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Project Plan',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('SDS|Software Design Specification',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('^SRS|^Software Requirement Specification|Software Requirements Specification',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Test Plan',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Release Notes',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Low Level Design',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('High Level Design',d,re.IGNORECASE):
    l.append(d)
    
   elif re.search('Configuration Management Plan',d,re.IGNORECASE):
    l.append(d)
   
 
  sow=['introduction','project overview and scope','definitions and acronyms','references','assumptions, dependencies, constraints and risks','high level requirements/target specifications','estimation summary','responsibility and ownership','key milestones and work products to be delivered','release definition','development and test environment','knowledge transfer','verification methods and acceptance criteria','tracking and reporting','product delivery mechanism','attachments']
  srs=['introduction','assumptions and risks','dependencies, constraints and limitations','overall description','functional requirements','interface requirements','non-functional requirements','error handling','compiler defined options','acceptance criteria','appendix']
  sds=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design - (modules 1..n)','sub-module 1..n design','sub-module 1..n pseudo code','integration approach and sequence','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  um=['introduction','environment setup','limitations and constraints','over all description','states of the system','[module_name]','source code build procedure','appendix i','trouble shooting']
  rn=['project details','contents, media and release method','brief introduction of product','summary of release history','environment required','limitations and known issues','installation procedure','components of release']
  tp=['introduction','definitions and acronyms','references','assumptions, dependencies and constraints','scope','test environment','features not tested']
  hld=['introduction','definitions and acronyms','references','design considerations','design strategy','system overview and software architecture','high level design – (modules 1..n)','integration approach and sequence:','logical data design/model','interface design','error handling and recovery','application library','build procedure/make file']
  lld=['introduction','definitions and acronyms','references','conventions and standards','data definition','sub-module 1..n design','sub-module 1..n pseudo code']
  cmp=['scope','acronyms and definitions','cm tools','project environment','configuration identification','naming conventions','location of cis','versioning','baselining','configuration control','configuration status accounting','configuration management audit','backup','release mechanism']
  pp=['introduction','definitions and acronyms','strategy and approach','assumptions, constraints and dependencies','project organization structure','risk management plan','summary of estimates by phase / delivery milestones','development environment','project monitoring and control /status reporting/ communication plan','quantitative objectives, measurement and data management plan','quality management (verification, validation and causal analysis)','product release plan','standards to be followed','decision management plan','tailoring and deviations','reference sheet_1ument']

   
  if re.search("user manual",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in um:
     if re.search(j,i):
      lll.append(j)			##lll is list of headings found in main document and are present in table of contents.
   result=set(um)-set(lll)
  elif re.search("Release Notes",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in rn:
     if re.search(j,i):
      lll.append(j)
   result=set(rn)-set(lll)
  elif re.search("SRS|Software Requirement Specification|Software Requirements Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in srs:
     if re.search(j,i):
      lll.append(j)
   print(lll)
   result=set(srs)-set(lll)
  elif re.search("Statement of Work|Scope of Work",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sow:
     if re.search(j,i):
      lll.append(j)
   result=set(sow)-set(lll)
  elif re.search("Project Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in pp:
     if re.search(j,i):
      lll.append(j)
   result=set(pp)-set(lll)
  elif re.search("^Test Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in tp:
     if re.search(j,i):
      lll.append(j)
   result=set(tp)-set(lll)
  elif re.search("Low Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in lld:
     if re.search(j,i):
      lll.append(j)
   result=set(lld)-set(lll)
  elif re.search("High Level Design",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in hld:
     if re.search(j,i):
      lll.append(j)
   result=set(hld)-set(lll)
  elif re.search("Configuration Management Plan",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in cmp:
     if re.search(j,i):
      lll.append(j)
   result=set(cmp)-set(lll)
  elif re.search("SDS|Software Design Specification",str(l[0]),re.IGNORECASE):
   for i in ll:
    for j in sds:
     if re.search(j,i):
      lll.append(j)
   result=set(sds)-set(lll)
  #print(result)
  if result==set():
   print("Document template is according to QMS.")
   print("Status:Pass")
  else:
   for x in result:
    print("Heading Not Found:","\t",x.upper()) 
   print("Status:Fail")  
end=datetime.now()
print("\nDocument Review End Time:", end)
print("\nTime taken For Document Review:", end-start,"HH:MM:SS")  